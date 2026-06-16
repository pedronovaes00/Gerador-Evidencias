#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Processador de template para DocJobs
Este módulo processa templates .docx substituindo campos específicos
"""

import os
import re
from datetime import datetime
from docx import Document
import json

class DocJobsTemplateProcessor:
    def __init__(self, template_path):
        self.template_path = template_path
        self.field_mappings = {
            # Mapeamento dos campos do front-end para os campos do template
            'smNumber': 'Número dda SM:',
            'callNumber': 'Numero do Chamado:',
            'requester': 'Solicitante:',
            'phone': 'Telefone:',
            'requestType': 'Tipo Solicitação:',
            'environment': 'Ambiente:',
            'requestDate': 'Data Solicitação:',
            'deadlineDate': 'Data limite p/ produção:',
            'subject': 'Assunto:',
            'businessUnit': 'Unidade de Negócio:',
            'system': 'Sistema:',
            'businessImpact': 'Impacto do Negócio:',
            'criticality': 'Criticidade:',
            'jobName': 'Nome do Job:',
            'owner': 'Owner:',
            'shellName': 'Nome do Shell:',
            'path': 'Path:',
            'server': 'Servidor:',
            'ip': 'IP:',
            'jobDescription': 'Descrição do Job:',
            'frequencyType': 'Frequência:',
            'startTime': 'Horário Início (hh:mm):',
            'endTime': 'Horário Final (hh:mm):',
            'frequencyDetails': 'Detalhamento da Frequência:',
            'predecessorProcess': 'Sub processo(s) Antecessor(es):',
            'successorProcess': 'Sub processo(s) Sucessor(es):',
            'parallelExecution': 'Execução Paralela?',
            'processNames': 'Nome dos Processos:',
            'param1': 'Parm1:',
            'param2': 'Parm2:',
            'param3': 'Parm3:',
            'param4': 'Parm4:',
            'param5': 'Parm5:',
            'param6': 'Parm6:',
            'param7': 'Parm7:',
            'param8': 'Parm8:',
            'returnCode': 'Return Code:',
            'emailNotification': 'E-mail destino para notificação:',
            'observations': 'Observação:',
            'executor': 'Executor:',
            'executionDate': 'Data de Execução:',
            'approver': 'Aprovador:',
            'approvalDate': 'Data de Aprovação:'
        }
    
    def extract_template_structure(self):
        """Extrai a estrutura do template para análise"""
        try:
            doc = Document(self.template_path)
            structure = {
                'paragraphs': [],
                'tables': []
            }
            
            # Extrair parágrafos
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    structure['paragraphs'].append({
                        'index': i,
                        'text': paragraph.text.strip()
                    })
            
            # Extrair tabelas
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    'table_index': table_idx,
                    'rows': []
                }
                
                for row_idx, row in enumerate(table.rows):
                    row_data = {
                        'row_index': row_idx,
                        'cells': []
                    }
                    
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip().replace('\u2002', ' ').replace('\n', ' ')
                        if cell_text:
                            row_data['cells'].append({
                                'cell_index': cell_idx,
                                'text': cell_text
                            })
                    
                    if row_data['cells']:
                        table_data['rows'].append(row_data)
                
                if table_data['rows']:
                    structure['tables'].append(table_data)
            
            return structure
        except Exception as e:
            print(f"Erro ao extrair estrutura do template: {str(e)}")
            return None
    
    def find_field_in_cell(self, cell_text, field_label):
        """Verifica se um campo está presente no texto da célula"""
        return field_label.lower() in cell_text.lower()
    
    def replace_field_value(self, cell_text, field_label, new_value):
        """Substitui o valor de um campo no texto da célula"""
        # Para campos de horário, manter o formato original completo
        if 'Horário' in field_label and '(hh:mm)' in field_label:
            if '(hh:mm):' in cell_text:
                # Encontrar a posição do último ':' para preservar '(hh:mm):'
                last_colon_pos = cell_text.rfind(':')
                if last_colon_pos != -1:
                    label_part = cell_text[:last_colon_pos + 1]  # Inclui o ':'
                    return f"{label_part} {new_value}"
            return f"{field_label} {new_value}"
        
        # Para outros campos, usar a lógica original
        if field_label.lower() in cell_text.lower():
            # Substituir mantendo o formato original
            if ':' in cell_text:
                parts = cell_text.split(':', 1)
                return f"{parts[0]}: {new_value}"
            else:
                return f"{field_label} {new_value}"
        
        return cell_text
    
    def process_template(self, form_data, output_path):
        """Processa o template substituindo os campos com os dados do formulário"""
        try:
            # Carregar o documento template
            doc = Document(self.template_path)
            
            print(f"Processando template: {self.template_path}")
            print(f"Dados recebidos: {json.dumps(form_data, indent=2, ensure_ascii=False)}")
            
            # Processar tabelas
            processed_fields = set()  # Rastrear campos já processados
            
            for table_idx, table in enumerate(doc.tables):
                print(f"\nProcessando tabela {table_idx}")
                
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        original_text = cell.text.strip()
                        if not original_text:
                            continue
                        
                        # Verificar cada campo do formulário
                        for form_field, form_value in form_data.items():
                            if form_field in self.field_mappings:
                                field_label = self.field_mappings[form_field]
                                
                                # Verificar se este campo já foi processado
                                if form_field in processed_fields:
                                    continue
                                
                                if self.find_field_in_cell(original_text, field_label):
                                    # Processar valores especiais
                                    processed_value = self.process_special_values(form_field, form_value, form_data)
                                    
                                    # Substituir o valor
                                    new_text = self.replace_field_value(original_text, field_label, processed_value)
                                    cell.text = new_text
                                    
                                    print(f"  Célula [{row_idx}][{cell_idx}]: '{original_text}' -> '{new_text}'")
                                    processed_fields.add(form_field)
                                    break
            
            # Configurar propriedades do documento para garantir editabilidade
            try:
                # Definir propriedades do documento
                doc.core_properties.title = "Documento de Job - Editável"
                doc.core_properties.author = "Sistema Gerador de Evidências"
                doc.core_properties.subject = "Documento gerado automaticamente"
                doc.core_properties.comments = "Documento editável gerado pelo sistema"
                doc.core_properties.created = datetime.now()
                doc.core_properties.modified = datetime.now()
                
                # Salvar o documento processado
                doc.save(output_path)
                print(f"\nDocumento salvo em: {output_path}")
                
                # Verificar se o arquivo foi salvo corretamente
                if os.path.exists(output_path):
                    file_size = os.path.getsize(output_path)
                    print(f"✅ Arquivo criado com sucesso - Tamanho: {file_size} bytes")
                    
                    # Tentar reabrir o arquivo para validar
                    try:
                        test_doc = Document(output_path)
                        print(f"✅ Arquivo validado - {len(test_doc.paragraphs)} parágrafos, {len(test_doc.tables)} tabelas")
                    except Exception as e:
                        print(f"⚠️ Aviso: Erro ao validar arquivo: {e}")
                else:
                    print("❌ Erro: Arquivo não foi criado")
                    return False
                    
            except Exception as save_error:
                print(f"❌ Erro ao salvar documento: {save_error}")
                return False
            
            return True
            
        except Exception as e:
            print(f"Erro ao processar template: {str(e)}")
            return False
    
    def process_special_values(self, field_name, field_value, all_data):
        """Processa valores especiais baseados no tipo de campo"""
        # Processar tipo de frequência
        if field_name == 'frequencyType':
            return field_value or 'Diário'
        
        # Processar horários
        if field_name in ['startTime', 'endTime']:
            if field_value:
                return field_value
            elif field_name == 'startTime':
                return '01:00'
            else:
                return ''
        
        # Processar dias da semana para frequencyDetails
        if field_name == 'frequencyDetails':
            # Se o usuário preencheu um valor personalizado, usar APENAS ele
            if field_value and field_value.strip():
                return field_value.strip()
            
            # Caso contrário, gerar automaticamente baseado no tipo e dias selecionados
            weekdays = []
            day_mapping = {
                'monday': 'Segunda',
                'tuesday': 'Terça', 
                'wednesday': 'Quarta',
                'thursday': 'Quinta',
                'friday': 'Sexta',
                'saturday': 'Sábado',
                'sunday': 'Domingo'
            }
            
            for day_key, day_name in day_mapping.items():
                if all_data.get(day_key) == 'on':
                    weekdays.append(day_name)
            
            # Obter o tipo de frequência
            frequency_type = all_data.get('frequencyType', 'Diário')
            
            if weekdays:
                if frequency_type == 'Mensal':
                    return f"Mensalmente - Dias: {', '.join(weekdays)}"
                else:
                    return f"Diariamente - Dias: {', '.join(weekdays)}"
            else:
                return frequency_type or 'Diário'
        
        # Processar observações
        if field_name == 'observations':
            if field_value and field_value.strip():
                return field_value.strip()
            else:
                return 'O Param3 DATA é opcional e deve ser preenchido em casos de execuções retroativas. <DATA>: O parâmetro de data tem formato AAAAMMDD.'
        
        # Processar valores booleanos
        if field_name == 'parallelExecution':
            return 'Sim' if field_value in ['on', 'Sim'] else 'Não'
        
        # Processar campos de data
        if 'Date' in field_name and not field_value:
            return datetime.now().strftime('%d/%m/%Y')
        
        # Processar campos vazios
        if not field_value or field_value.strip() == '':
            return 'N/A'
        
        return str(field_value)
    
    def validate_required_fields(self, form_data):
        """Valida se todos os campos obrigatórios estão preenchidos"""
        required_fields = {
            'smNumber': 'Número da SM',
            'callNumber': 'Número do Chamado',
            'requester': 'Solicitante',
            'phone': 'Telefone',
            'requestDate': 'Data Solicitação',
            'deadlineDate': 'Data limite p/ produção',
            'subject': 'Assunto',
            'system': 'Sistema',
            'businessImpact': 'Impacto do Negócio',
            'criticality': 'Criticidade',
            'jobName': 'Nome do Job',
            'shellName': 'Nome do Shell',
            'path': 'Path',
            'server': 'Servidor',
            'ip': 'IP',
            'jobDescription': 'Descrição do Job',
            'returnCode': 'Return Code'
        }
        
        missing_fields = []
        for field_id, field_label in required_fields.items():
            value = form_data.get(field_id)
            if not value or str(value).strip() == '':
                missing_fields.append({
                    'field': field_id,
                    'label': field_label
                })
        
        # Validação especial para frequencyDetails - pode ser gerado automaticamente
        frequency_details = form_data.get('frequencyDetails', '').strip()
        if not frequency_details:
            # Verificar se temos dados suficientes para gerar automaticamente
            frequency_type = form_data.get('frequencyType', 'Diário')
            has_weekdays = any(form_data.get(day) == 'on' for day in 
                             ['monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday'])
            
            if not has_weekdays and not frequency_details:
                missing_fields.append({
                    'field': 'frequencyDetails',
                    'label': 'Detalhamento da Frequência (ou selecione pelo menos um dia da semana)'
                })
        
        return missing_fields

def test_template_processor():
    """Função de teste para o processador de template"""
    # Dados de teste
    test_data = {
        'smNumber': 'SM-12345',
        'callNumber': 'CH-67890',
        'requester': 'João Silva',
        'phone': '(11) 99999-9999',
        'requestType': 'Criação de Job',
        'environment': 'Produção',
        'requestDate': '10/01/2025',
        'deadlineDate': '15/01/2025',
        'subject': 'Teste de processamento de template',
        'businessUnit': 'DRDR - Tarifação',
        'system': 'BRM',
        'businessImpact': 'Alto impacto no negócio',
        'criticality': 'Alta',
        'jobName': 'teste_processamento_template',
        'shellName': 'exec_teste.sh',
        'path': '/teste/scripts/',
        'server': 'SERVIDOR_TESTE',
        'ip': '192.168.1.100',
        'jobDescription': 'Job de teste para validação do processamento',
        'frequencyDetails': 'Diário',
        'monday': 'on',
        'tuesday': 'on',
        'wednesday': 'on',
        'thursday': 'on',
        'friday': 'on',
        'predecessorProcess': 'N/A',
        'successorProcess': 'N/A',
        'parallelExecution': 'Não',
        'processNames': 'N/A',
        'param1': '/teste/param1',
        'param2': 'TESTE_PARAM2',
        'param3': '<DATA>',
        'param4': '',
        'param5': '',
        'param6': '',
        'param7': '',
        'param8': '',
        'returnCode': 'Exit 0 = sucesso, Exit != 0 = erro',
        'emailNotification': 'teste@empresa.com',
        'executor': 'Testador',
        'executionDate': '10/01/2025',
        'approver': 'Aprovador Teste',
        'approvalDate': '10/01/2025'
    }
    
    # Caminhos
    template_path = 'docjobs/c_alarmes_valida_selecao_arrecadado.docx'
    output_path = 'docjobs/gerados/teste_processamento_template.docx'
    
    # Criar diretório de saída se não existir
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Processar template
    processor = DocJobsTemplateProcessor(template_path)
    
    # Validar campos obrigatórios
    missing_fields = processor.validate_required_fields(test_data)
    if missing_fields:
        print(f"Campos obrigatórios não preenchidos: {missing_fields}")
        return False
    
    # Processar o template
    success = processor.process_template(test_data, output_path)
    
    if success:
        print(f"\n✅ Template processado com sucesso!")
        print(f"📄 Arquivo gerado: {output_path}")
        return True
    else:
        print(f"\n❌ Erro ao processar template")
        return False

if __name__ == '__main__':
    test_template_processor()